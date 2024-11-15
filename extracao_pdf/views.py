from django.shortcuts import render, redirect, get_object_or_404
from django.core.files.storage import FileSystemStorage
from .forms import UploadArquivoForm
from .utils import extrair_texto_com_orientacao, encontrar_numero_processo, get_filtered_processos
from .models import (Processo, MyModel, Leiloeiro, Vara, TipoDocumento, Justica, Board, FrasePadrao, Estado, \
    List, Card, Comment, GrupoResponsavel, ProcessoArquivo, EmailLog, EmailTemplate, ChecklistItem, Cidade,
                     CardsAgregados)
from django.template import Context, Template
from django.http import HttpResponse
from docxtpl import DocxTemplate
from django.contrib.staticfiles import finders
from bs4 import BeautifulSoup, NavigableString, Tag
import pypandoc
import os
from django.conf import settings
import requests
from rest_framework import viewsets
from extracao_pdf.serializer import ProcessoSerializer, BoardSerializer
from io import BytesIO
import html2text
from datetime import datetime, timedelta
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
import pytz
from django import forms
from django.contrib import messages
from django.db.models import Q, Min, F, Count, Prefetch
from django.shortcuts import render
from collections import defaultdict
from django.http import JsonResponse
from .models import Leilao, DataLeilao, Board
from django.utils.timezone import localtime, now, make_naive
import logging
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from .utils import enviar_notificacao
from django.views.decorators.http import require_POST
from django.views.decorators.csrf import csrf_exempt
import json
from notifications.models import Notification
from django.core.paginator import Paginator
import re
from notifications.signals import notify
from django.views import View
from django.dispatch import receiver
from django.db.models.signals import post_save
from django.contrib.admin import SimpleListFilter
from django.core.mail import send_mail
from django.utils import timezone
from django.utils.dateformat import format as date_format
import pytz
from num2words import num2words
import locale
from babel.numbers import format_currency
from decimal import Decimal, ROUND_HALF_UP
from dal import autocomplete
from django.template.loader import render_to_string
from weasyprint import HTML, CSS
from django.contrib.admin.views.decorators import staff_member_required
from docx.shared import Pt
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT


locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')

logger = logging.getLogger(__name__)


class BoardViewSet(viewsets.ModelViewSet):
    queryset = Board.objects.all()
    serializer_class = BoardSerializer


def admin_relatorios(request):
    # Inicialmente, nenhum processo é listado
    processos = Processo.objects.none()

    estado = request.GET.get('estado')
    leiloeiro = request.GET.get('leiloeiro')
    dias_antes_leilao = request.GET.get('dias_antes_leilao')
    dias_depois_cadastro = request.GET.get('dias_depois_cadastro')
    intervalo_inicio = request.GET.get('intervalo_inicio')
    intervalo_fim = request.GET.get('intervalo_fim')
    coluna_kanban = request.GET.get('coluna_kanban')

    estados = Estado.objects.all()
    leiloeiros = Leiloeiro.objects.all()
    colunas_kanban = List.objects.all()

    if estado or leiloeiro or dias_antes_leilao or dias_depois_cadastro or (intervalo_inicio and intervalo_fim) or coluna_kanban:
        processos = get_filtered_processos(request)

    context = {
        'processos': processos,
        'estados': estados,
        'leiloeiros': leiloeiros,
        'colunas_kanban': colunas_kanban,
    }

    return render(request, 'extracao_pdf/admin_relatorios.html', context)


def kanban_board(request, board_id):
    board = get_object_or_404(Board, pk=board_id)

    lists = board.lists.all().order_by('ordem')

    for list_ in lists:
        list_.ordered_cards = list_.cards.annotate(
            first_leilao_data=Min('processo__leilao__datas_leilao__data_hora')
        ).order_by('first_leilao_data')

    grupos_responsaveis = GrupoResponsavel.objects.all()

    responsaveis = User.objects.filter(cards_responsaveis__list__board=board).distinct()

    return render(request, 'kanban.html', {
        'board': board,
        'lists': lists,
        'grupos_responsaveis': grupos_responsaveis,
        'responsaveis': responsaveis,
    })


def add_list(board_id, name, position):
    board = Board.objects.get(id=board_id)
    existing_lists = board.lists.order_by('order')

    # Atualiza a ordem das listas existentes
    for lst in existing_lists:
        if lst.order >= position:
            lst.order += 1
            lst.save()

    # Cria a nova lista na posição desejada
    new_list = List.objects.create(name=name, board=board, order=position)
    return new_list


def notifications_view(request):
    user = request.user
    notifications_list = Notification.objects.filter(recipient=user).order_by('-timestamp')
    paginator = Paginator(notifications_list, 20)  # 20 notificações por página

    page_number = request.GET.get('page')
    notifications = paginator.get_page(page_number)

    notifications_data = []
    for notification in notifications:
        card_data = None
        if notification.target_object_id:
            try:
                card = Card.objects.get(id=notification.target_object_id)
                board_id = card.list.board.id
                card_data = {
                    'board_id': board_id,
                    'card_id': card.id,
                }
            except Card.DoesNotExist:
                card_data = None

        notifications_data.append({
            'notification': notification,
            'card_data': card_data,
        })

    return render(request, 'notifications.html',
                  {'notifications': notifications, 'notifications_data': notifications_data})

def mark_as_read(request, notification_id):
    if request.method == 'POST':
        notification = get_object_or_404(Notification, id=notification_id, recipient=request.user)
        notification.unread = False
        notification.save()
    return redirect('notifications')


def calendario(request):
    return render(request, 'extracao_pdf/calendario.html')

def kanban_list(request):
    return render(request, 'extracao_pdf/kanban_list.html')



def api_leiloes(request):
    leiloes = Leilao.objects.all()
    eventos_agrupados = defaultdict(list)

    for leilao in leiloes:
        for data_leilao in leilao.datas_leilao.all():
            data_key = (leilao.id, localtime(data_leilao.data_hora).date())
            description = f"{data_leilao.etapa} - {data_leilao.tipo_data} {data_leilao.complemento} {localtime(data_leilao.data_hora).strftime('%d/%m/%Y %H:%M')}"
            eventos_agrupados[data_key].append({
                'title': leilao.nome,
                'start': data_leilao.data_hora,
                'description': description,
                'time': localtime(data_leilao.data_hora).strftime('%H:%M')
            })

    eventos_consolidados = []
    for (leilao_id, data), eventos_dia in eventos_agrupados.items():
        if len(eventos_dia) > 1:
            descriptions = "; ".join([f"{evento['description']}" for evento in eventos_dia])
            eventos_consolidados.append({
                'title': eventos_dia[0]['title'],
                'start': data.isoformat(),
                'description': descriptions,
            })
        else:
            evento = eventos_dia[0]
            eventos_consolidados.append({
                'title': evento['title'],
                'start': data.isoformat() if evento['time'] == "00:00" else evento['start'].isoformat(),
                'description': evento['description'],
            })

    return JsonResponse(eventos_consolidados, safe=False)


class SimpleUploadForm(forms.Form):
    arquivo = forms.FileField()

def upload_pdf(request):
    if request.method == 'POST':
        form = SimpleUploadForm(request.POST, request.FILES)
        if form.is_valid():
            arquivo = request.FILES['arquivo']
            fs = FileSystemStorage()
            filename = fs.save(arquivo.name, arquivo)
            arquivo_url = fs.url(filename)

            try:
                texto_extraido = extrair_texto_com_orientacao(arquivo)
                numero_processo = encontrar_numero_processo(" ".join(texto_extraido))
                request.session['texto_extraido'] = texto_extraido
                request.session['numero_processo'] = numero_processo
                request.session['arquivo_pdf_url'] = arquivo_url

                # Não deletar o arquivo aqui; necessário para visualização
                # fs.delete(filename)

                return redirect('pre_visualizacao')
            except Exception as e:
                messages.error(request, f"Erro ao processar o arquivo: {str(e)}")
                # Deleta o arquivo se ocorrer erro no processamento
                os.remove(os.path.join(fs.location, filename))
    else:
        form = SimpleUploadForm()

    return render(request, 'extracao_pdf/upload_pdf.html', {'form': form})

def pre_visualizacao(request):
    numero_processo = request.session.get('numero_processo')
    arquivo_pdf_url = request.session.get('arquivo_pdf_url')
    texto_extraido = request.session.get('texto_extraido')

    if not numero_processo or not arquivo_pdf_url or not texto_extraido:
        messages.error(request, "Informações necessárias para visualização não estão disponíveis. Por favor, faça o upload novamente.")
        return redirect('upload_pdf')

    return render(request, 'extracao_pdf/resultados.html', {
        'numero_processo': numero_processo,
        'arquivo_pdf_url': arquivo_pdf_url,
        'texto_extraido': texto_extraido
    })


def confirmar_salvar_processo(request):
    numero_processo = request.session.get('numero_processo', None)

    if numero_processo:
        response = redirect(f'/admin/extracao_pdf/processo/add/?numero_processo={numero_processo}')

        if 'numero_processo' in request.session:
            del request.session['numero_processo']
        if 'texto_extraido' in request.session:
            del request.session['texto_extraido']
        if 'arquivo_pdf_url' in request.session:
            del request.session['arquivo_pdf_url']

        return response
    else:
        return redirect('upload_pdf')


def escolher_documento(request, processo_id):
    processo = get_object_or_404(Processo, pk=processo_id)
    vara = processo.vara
    justica = vara.justica if vara else None

    modelos = MyModel.objects.filter(
        Q(vara__isnull=True) | Q(vara=vara),
        Q(justica__isnull=True) | Q(justica=justica),
        Q(leiloeiro__isnull=True) | Q(leiloeiro=processo.leiloeiro)
    )

    return render(request, 'admin/escolher_documento.html', {
        'processo': processo,
        'modelos': modelos
    })


def gerar_documento(request, processo_id, tipo_documento_id):
    processo = get_object_or_404(Processo, pk=processo_id)
    tipo_documento = get_object_or_404(TipoDocumento, pk=tipo_documento_id)
    modelo = MyModel.objects.filter(tipo_documento=tipo_documento).first()

    if modelo is None:
        return HttpResponse("Modelo de documento não encontrado.", status=404)

    conteudo_html = modelo.content.replace('{{numero}}', str(processo.numero))
    print("Conteúdo HTML antes da conversão:", conteudo_html)

    h = html2text.HTML2Text()
    h.ignore_links = True
    texto_documento = h.handle(conteudo_html)
    print("Conteúdo HTML após a conversão para texto:", texto_documento)

    template_path = finders.find('modelos/template.docx')
    if not template_path:
        return HttpResponse("Arquivo de template não encontrado.", status=404)

    doc = DocxTemplate(template_path)

    context = {'conteudo': texto_documento}

    doc.render(context)

    f = BytesIO()
    doc.save(f)
    f.seek(0)

    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{tipo_documento}_{processo.pk}.docx"'

    return response


def format_date(date):
    months = [
        "janeiro", "fevereiro", "março", "abril", "maio", "junho",
        "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"
    ]
    return f"{date.day} de {months[date.month - 1]} de {date.year}"


def clean_html(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')

    for p in soup.find_all('p'):
        while p.contents and isinstance(p.contents[0], Tag) and p.contents[0].name == 'br':
            p.contents[0].extract()

    return str(soup)


def valor_por_extenso(valor):
    valor_inteiro = int(valor)
    centavos = int(round((valor - valor_inteiro) * 100))
    valor_extenso = num2words(valor_inteiro, lang='pt_BR').replace(' e ', ' ').replace('-', ' ')
    if centavos > 0:
        centavos_extenso = num2words(centavos, lang='pt_BR').replace(' e ', ' ').replace('-', ' ')
        return f"{valor_extenso} reais e {centavos_extenso} centavos"
    return f"{valor_extenso} reais"


def gerar_texto_executados(executados):
    executados_texto = []
    for executado in executados:
        if executado.tipo_pessoa == 'PF':
            texto = f"{executado.nome} e seu(a) cônjuge se casado(a) for"
        elif executado.tipo_pessoa == 'PJ':
            texto = f"{executado.nome} na pessoa de seu(s) representante(s) legal(is)"
        executados_texto.append(texto)
    return '; '.join(executados_texto)


from docx.oxml.ns import qn


def replace_placeholder_in_paragraph(paragraph, placeholder, new_text, bold=False):
    # Concatena todo o texto do parágrafo em uma única string
    full_text = ''.join(run.text for run in paragraph.runs)

    if placeholder in full_text:
        original_font_size = paragraph.runs[0].font.size if paragraph.runs else None

        full_text = full_text.replace(placeholder, new_text)

        for run in paragraph.runs:
            run.clear()

        run = paragraph.add_run(full_text)
        if original_font_size:
            run.font.size = original_font_size

        if bold:
            run.bold = True


def gerar_tabela_dinamica(doc, elem):
    num_rows = len(elem.find_all('tr'))
    num_cols = len(elem.find_all('tr')[0].find_all(['td', 'th'])) if num_rows > 0 else 0

    if num_rows == 0 or num_cols == 0:
        return

    tabela_docx = doc.add_table(rows=num_rows, cols=num_cols)

    try:
        tabela_docx.style = 'Table Grid'
    except KeyError:
        pass

    for row_idx, row in enumerate(elem.find_all('tr')):
        cells = row.find_all(['td', 'th'])
        for col_idx, cell in enumerate(cells):
            cell_docx = tabela_docx.cell(row_idx, col_idx)
            cell_content = cell.contents

            paragraph = cell_docx.paragraphs[0]
            paragraph.clear()

            for content in cell_content:
                if isinstance(content, Tag) and content.name == 'strong':
                    run = paragraph.add_run(content.get_text(strip=True))
                    run.bold = True
                elif isinstance(content, NavigableString):
                    run = paragraph.add_run(content)

def gerar_docx_a_partir_de_html(request, modelo_id, processo_id):
    modelo = get_object_or_404(MyModel, pk=modelo_id)
    processo = get_object_or_404(Processo, pk=processo_id)

    documento_html = clean_html(modelo.content)

    venda_direta_marcada = processo.venda_direta == 1

    frases_padrao = FrasePadrao.objects.all()
    for frase in frases_padrao:
        if frase.chave == 'vendadireta_sugestao':
            if venda_direta_marcada:
                documento_html = documento_html.replace(f'{{{{{frase.chave}}}}}', frase.descricao)
            else:
                documento_html = documento_html.replace(f'{{{{{frase.chave}}}}}', '')
        elif frase.chave in ['frase_veiculo', 'frase_imovel']:
            documento_html = documento_html.replace(f'{{{{{frase.chave}}}}}', frase.descricao)


    documento_html = documento_html.replace('{{numero_processo}}', str(processo.numero))
    depositarios = '; '.join([
        f"<strong>{str(index + 1).zfill(2)})</strong> {str(bem.depositario)}"
        for index, bem in enumerate(processo.bens.all())
        if bem.depositario
    ])
    documento_html = documento_html.replace('{{depositario}}', depositarios)
    documento_html = documento_html.replace('{{onus}}', str(processo.onus or ''))
    documento_html = documento_html.replace('{{fls_divida}}',
                                            str(processo.fls_divida) if processo.fls_divida else 'Folha não especificada')
    exequentes_nomes = ', '.join([exequente.nome for exequente in processo.exequentes.all()])
    documento_html = documento_html.replace('{{nome_exequente}}', exequentes_nomes)
    executados_nomes = ', '.join([executado.nome for executado in processo.executados.all()])
    documento_html = documento_html.replace('{{nome_executado}}', executados_nomes)
    texto_executados = gerar_texto_executados(processo.executados.all())
    documento_html = documento_html.replace('{{executados}}', texto_executados)
    documento_html = documento_html.replace('{{email_vara}}', processo.vara.email if processo.vara else 'Sem Vara')
    documento_html = documento_html.replace('{{telefone_vara}}', processo.vara.telefone if processo.vara else 'Sem Vara')
    documento_html = documento_html.replace('{{endereco_completo_vara}}', processo.vara.endereco_completo() if processo.vara else 'Sem Vara')
    endereco_vara_maiusculo = processo.vara.endereco_completo().upper() if processo.vara else 'SEM VARA'
    documento_html = documento_html.replace('{{endereco_completo_vara_maiusculo}}', endereco_vara_maiusculo)
    documento_html = documento_html.replace('{{nome_vara}}', processo.vara.nome if processo.vara else 'Sem Vara')
    documento_html = documento_html.replace('{{juiz}}', processo.vara.juiz if processo.vara else 'Sem Juiz')
    juiz = processo.vara.juiz if processo.vara else 'Sem Juiz'
    juiz_maiusculo = juiz.upper()
    documento_html = documento_html.replace('{{juiz_maiusculo}}', juiz_maiusculo)

    if processo.vara and processo.vara.justica:
        tipo_justica = processo.vara.justica.tipo
        if tipo_justica == 'justica_estadual':
            titulo_juiz = 'Juiz de Direito'
        elif tipo_justica == 'justica_federal':
            titulo_juiz = 'Juiz Federal'
        elif tipo_justica == 'justica_trabalho':
            titulo_juiz = 'Juiz do Trabalho'
        elif tipo_justica == 'justica_eleitoral':
            titulo_juiz = 'Juiz Eleitoral'
        else:
            titulo_juiz = 'Juiz'
    else:
        titulo_juiz = 'Sem Justiça Definida'

    documento_html = documento_html.replace('{{titulo_juiz}}', titulo_juiz)
    documento_html = documento_html.replace('{{titulo_juiz_maiusculo}}', titulo_juiz.upper())

    documento_html = documento_html.replace('{{nome_vara_maiusculas}}', processo.vara.nome.upper() if processo.vara else 'SEM VARA')
    documento_html = documento_html.replace('{{email_leiloeiro}}', processo.leiloeiro.email if processo.leiloeiro else 'Sem Leiloeiro')
    documento_html = documento_html.replace('{{nome_leiloeiro}}', processo.leiloeiro.nome if processo.leiloeiro else 'Sem Leiloeiro')
    documento_html = documento_html.replace('{{nome_leiloeiro_maiusculo}}', processo.leiloeiro.nome.upper() if processo.leiloeiro else 'SEM LEILOEIRO')
    documento_html = documento_html.replace('{{classe}}', processo.classe_processual.nome if processo.classe_processual else 'Sem Classe Processual')
    documento_html = documento_html.replace('{{cidade_vara}}', processo.vara.comarca.nome if processo.vara and processo.vara.comarca else 'Sem Cidade')
    documento_html = documento_html.replace('{{cidade_vara_maiusculo}}', processo.vara.comarca.nome.upper() if processo.vara and processo.vara.comarca else 'SEM CIDADE')
    documento_html = documento_html.replace('{{justica}}', processo.vara.justica.nome if processo.vara and processo.vara.justica else 'Sem Justiça')
    documento_html = documento_html.replace('{{justica_maiusculo}}', processo.vara.justica.nome.upper() if processo.vara and processo.vara.justica else 'SEM JUSTIÇA')
    uf = processo.vara.comarca.estado.uf if processo.vara and processo.vara.comarca and processo.vara.comarca.estado else 'UF não definida'
    documento_html = documento_html.replace('{{uf}}', uf)
    documento_html = documento_html.replace(
        '{{estado_completo_vara}}',
        f"{processo.vara.comarca.estado.artigo_definido} {processo.vara.comarca.estado.nome}"
        if processo.vara and processo.vara.comarca and processo.vara.comarca.estado
        else 'Estado não definido'
    )

    estado_completo_maiusculo = (
        f"{processo.vara.comarca.estado.artigo_definido} {processo.vara.comarca.estado.nome}".upper()
        if processo.vara and processo.vara.comarca and processo.vara.comarca.estado
        else 'ESTADO NÃO DEFINIDO'
    )

    documento_html = documento_html.replace('{{estado_completo_vara_maiusculo}}', estado_completo_maiusculo)

    documento_html = documento_html.replace('{{site}}', processo.leiloeiro.site if processo.leiloeiro else 'Sem Site')
    data_hoje = date_format(datetime.now(), 'd \\d\\e F \\d\\e Y')
    documento_html = documento_html.replace('{{data_hoje}}', data_hoje)

    comissao = processo.comissao
    if comissao == comissao.to_integral_value():
        comissao_percentual = f"{int(comissao)}%"
    else:
        comissao_percentual = f"{comissao:.2f}%"

    comissao_extenso = num2words(comissao, lang='pt_BR').replace('vírgula', 'virgula')
    comissao_placeholder = f"{comissao_percentual} ({comissao_extenso} por cento)"
    documento_html = documento_html.replace('{{comissao}}', comissao_placeholder)

    if processo.preco_vil is not None:
        if processo.preco_vil == int(processo.preco_vil):
            preco_vil_percentual = f"{int(processo.preco_vil)}%"
        else:
            preco_vil_percentual = f"{processo.preco_vil:.2f}%"

        preco_vil_extenso = num2words(processo.preco_vil, lang='pt_BR').replace('vírgula', 'virgula')
        preco_vil_placeholder = f"{preco_vil_percentual} ({preco_vil_extenso} por cento)"
    else:
        preco_vil_placeholder = "Preço vil não especificado"

    documento_html = documento_html.replace('{{preco_vil}}', preco_vil_placeholder)

    bens = processo.bens.all()
    bens_descricao = []
    bens_somente_descricao = []
    total_avaliacao = Decimal(0)

    for index, bem in enumerate(bens, start=1):
        valor_formatado = format_currency(bem.valor, 'BRL', locale='pt_BR').replace(u'\xa0', u' ')
        valor_extenso = valor_por_extenso(bem.valor)
        data_avaliacao_formatada = bem.data_avaliacao.strftime(
            '%d de %B de %Y') if bem.data_avaliacao else "data de avaliação não disponível"

        if bem.valor_atualizado and bem.data_avaliacao_atualizada:
            valor_atualizado_formatado = format_currency(bem.valor_atualizado, 'BRL', locale='pt_BR').replace(u'\xa0',
                                                                                                              u' ')
            valor_atualizado_extenso = valor_por_extenso(bem.valor_atualizado)
            data_atualizacao_formatada = bem.data_avaliacao_atualizada.strftime('%d de %B de %Y')

            descricao_bem = (
                f"{bem.descricao}, avaliado em {valor_formatado} ({valor_extenso}), em {data_avaliacao_formatada}. "
                f"Atualizado para {valor_atualizado_formatado} ({valor_atualizado_extenso}), em {data_atualizacao_formatada}")
        else:
            # Formato sem atualização
            descricao_bem = f"{bem.descricao}, avaliado em {valor_formatado} ({valor_extenso}), em {data_avaliacao_formatada}"

        if bens.count() == 1:
            descricao_bem = bem.descricao

        elif bens.count() > 1:
            descricao_bem = f"<strong>{index:02d})</strong> {descricao_bem}"
            if index < len(bens):
                descricao_bem += ";"

        bens_descricao.append(descricao_bem)
        bens_somente_descricao.append(bem.descricao)
        total_avaliacao += bem.valor or Decimal(0)

    documento_html = documento_html.replace('{{descricao_bens}}', ' '.join(bens_descricao))

    documento_html = documento_html.replace('{{descricao_somente_bens}}', ', '.join(bens_somente_descricao))

    # Placeholder para avaliação total dos bens
    total_avaliacao_formatado = format_currency(total_avaliacao, 'BRL', locale='pt_BR').replace(u'\xa0', u' ')
    total_avaliacao_extenso = valor_por_extenso(total_avaliacao)
    documento_html = documento_html.replace('{{avaliacao_total}}',
                                            f"{total_avaliacao_formatado} ({total_avaliacao_extenso}).")

    if bens.count() == 1:
        # Quando houver apenas um bem
        bem = bens.first()
        valor_formatado = format_currency(bem.valor, 'BRL', locale='pt_BR').replace(u'\xa0', u' ')
        valor_extenso = valor_por_extenso(bem.valor)
        data_avaliacao_formatada = bem.data_avaliacao.strftime(
            '%d de %B de %Y') if bem.data_avaliacao else "data de avaliação não disponível"

        # Verificar se há avaliação atualizada
        if bem.valor_atualizado and bem.data_avaliacao_atualizada:
            valor_atualizado_formatado = format_currency(bem.valor_atualizado, 'BRL', locale='pt_BR').replace(u'\xa0',
                                                                                                              u' ')
            valor_atualizado_extenso = valor_por_extenso(bem.valor_atualizado)
            data_atualizacao_formatada = bem.data_avaliacao_atualizada.strftime('%d de %B de %Y')

            # Formato com avaliação e atualização
            avaliacao_placeholder = (f"{valor_formatado} ({valor_extenso}), em {data_avaliacao_formatada}. "
                                     f"Atualizado para {valor_atualizado_formatado} ({valor_atualizado_extenso}), em {data_atualizacao_formatada}.")
        else:
            # Formato sem atualização
            avaliacao_placeholder = f"{valor_formatado} ({valor_extenso}), em {data_avaliacao_formatada}."
    else:
        # Quando houver mais de um bem - exibir o total
        avaliacao_placeholder = f"{total_avaliacao_formatado} ({total_avaliacao_extenso})."

    documento_html = documento_html.replace('{{avaliacao_dos_bens}}', avaliacao_placeholder)

    if processo.preco_vil is not None and total_avaliacao is not None:
        lance_minimo = total_avaliacao * (processo.preco_vil / 100)
        lance_minimo_formatado = format_currency(lance_minimo, 'BRL', locale='pt_BR').replace(u'\xa0', u' ')
        lance_minimo_extenso = valor_por_extenso(lance_minimo)
        lance_minimo_placeholder = f"{lance_minimo_formatado} ({lance_minimo_extenso})"
    else:
        lance_minimo_placeholder = "Lance mínimo não especificado"

    documento_html = documento_html.replace('{{lance_minimo}}', lance_minimo_placeholder)

    if processo.valor_divida and processo.data_divida:
        valor_divida_extenso = valor_por_extenso(processo.valor_divida)
        data_divida_formatada = processo.data_divida.strftime('%d de %B de %Y')

        valor_divida_formatado = f"{processo.valor_divida:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

        documento_html = documento_html.replace(
            '{{valor_divida}}',
            f"R$ {valor_divida_formatado} ({valor_divida_extenso}), em {data_divida_formatada}"
        )
    else:
        documento_html = documento_html.replace('{{valor_divida}}', 'Valor da dívida não especificado')

    bens_localizacao = []
    for index, bem in enumerate(bens, start=1):
        localizacao_texto = f"{index:02d}) {bem.localizacao if bem.localizacao else 'Não informado'}."
        bens_localizacao.append(localizacao_texto)

    # Se houver mais de um bem, use a formatação numerada, caso contrário, apenas a localização simples
    if len(bens) > 1:
        localizacao_final = ' '.join(bens_localizacao)
    else:
        localizacao_final = bens_localizacao[0] if bens_localizacao else 'Não informado'

    documento_html = documento_html.replace('{{localizacao_bens}}', localizacao_final)

    if processo.leiloeiro:
        if processo.leiloeiro.sexo == 'F':
            titulo_leiloeiro = 'Leiloeira'
            artigo_leiloeiro = 'a Leiloeira'
            pronome_leiloeiro_capitalizado = 'Esta Leiloeira'
            pronome_leiloeiro_min = 'esta leiloeira'
            artigo_simples = 'a'
            deste_leiloeiro = 'desta Leiloeira'
            artigo_a_ao = 'à'
        else:
            titulo_leiloeiro = 'Leiloeiro'
            artigo_leiloeiro = 'o Leiloeiro'
            pronome_leiloeiro_capitalizado = 'Este Leiloeiro'
            pronome_leiloeiro_min = 'este leiloeiro'
            artigo_simples = 'o'
            deste_leiloeiro = 'deste Leiloeiro'
            artigo_a_ao = 'ao'
    else:
        titulo_leiloeiro = 'Sem Leiloeiro'
        artigo_leiloeiro = 'Sem Leiloeiro'
        pronome_leiloeiro_capitalizado = 'Sem Leiloeiro'
        pronome_leiloeiro_min = 'sem leiloeiro'
        artigo_simples = 'sem'
        deste_leiloeiro = 'sem leiloeiro'
        artigo_a_ao = 'sem'

    # Variáveis maiúsculas
    titulo_leiloeiro_maiusculo = titulo_leiloeiro.upper()
    artigo_leiloeiro_maiusculo = artigo_leiloeiro.upper()
    pronome_leiloeiro_capitalizado_maiusculo = pronome_leiloeiro_capitalizado.upper()
    pronome_leiloeiro_min_maiusculo = pronome_leiloeiro_min.upper()
    artigo_simples_maiusculo = artigo_simples.upper()
    deste_leiloeiro_maiusculo = deste_leiloeiro.upper()
    artigo_a_ao_maiusculo = artigo_a_ao.upper()

    # Substituições no documento
    documento_html = documento_html.replace('{{titulo_leiloeiro}}', titulo_leiloeiro)
    documento_html = documento_html.replace('{{artigo_leiloeiro}}', artigo_leiloeiro)
    documento_html = documento_html.replace('{{pronome_leiloeiro}}', pronome_leiloeiro_capitalizado)
    documento_html = documento_html.replace('{{pronome_leiloeiro_min}}', pronome_leiloeiro_min)
    documento_html = documento_html.replace('{{artigo_simples}}', artigo_simples)
    documento_html = documento_html.replace('{{deste_leiloeiro}}', deste_leiloeiro)
    documento_html = documento_html.replace('{{artigo_a_ao_leiloeiro}}',
                                            artigo_a_ao)

    # Substituições maiúsculas no documento
    documento_html = documento_html.replace('{{titulo_leiloeiro_maiusculo}}', titulo_leiloeiro_maiusculo)
    documento_html = documento_html.replace('{{artigo_leiloeiro_maiusculo}}', artigo_leiloeiro_maiusculo)
    documento_html = documento_html.replace('{{pronome_leiloeiro_maiusculo}}', pronome_leiloeiro_capitalizado_maiusculo)
    documento_html = documento_html.replace('{{pronome_leiloeiro_min_maiusculo}}', pronome_leiloeiro_min_maiusculo)
    documento_html = documento_html.replace('{{artigo_simples_maiusculo}}',
                                            artigo_simples_maiusculo)
    documento_html = documento_html.replace('{{deste_leiloeiro_maiusculo}}', deste_leiloeiro_maiusculo)
    documento_html = documento_html.replace('{{artigo_a_ao_leiloeiro_maiusculo}}',
                                            artigo_a_ao_maiusculo)

    leiloeiro = processo.leiloeiro
    if leiloeiro:
        matricula = leiloeiro.matriculas.filter(
            estado=processo.vara.comarca.estado).first() if processo.vara and processo.vara.comarca else None
        if matricula:
            documento_html = documento_html.replace('{{numero_matricula}}', matricula.matricula)
            documento_html = documento_html.replace('{{cidade_leiloeiro}}',
                                                    matricula.cidade.nome if matricula.cidade else 'Cidade não cadastrada')

            endereco_completo = f"Rua {matricula.rua}, nº {matricula.numero}, {matricula.bairro}, CEP: {matricula.cep} - {matricula.cidade.nome}/{matricula.estado.uf}" if matricula.cidade and matricula.rua else "Endereço não cadastrado"
            documento_html = documento_html.replace('{{endereco_matricula}}', endereco_completo)
        else:
            documento_html = documento_html.replace('{{numero_matricula}}', 'Matrícula não encontrada')
            documento_html = documento_html.replace('{{cidade_leiloeiro}}', 'Matrícula não encontrada')
            documento_html = documento_html.replace('{{endereco_matricula}}', 'Endereço não disponível')
    else:
        documento_html = documento_html.replace('{{numero_matricula}}', 'Leiloeiro não encontrado')
        documento_html = documento_html.replace('{{cidade_leiloeiro}}', 'Leiloeiro não encontrado')
        documento_html = documento_html.replace('{{endereco_matricula}}', 'Leiloeiro não encontrado')

    local_tz = pytz.timezone('America/Sao_Paulo')

    datas_leilao = processo.leilao.datas_leilao.exclude(tipo_data='abertura') if processo.leilao else []

    datas_formatadas = []

    local_tz = pytz.timezone('America/Sao_Paulo')

    for data in datas_leilao:
        data_hora_local = data.data_hora.astimezone(local_tz)
        data_formatada = data_hora_local.strftime('%d/%m/%Y')
        datas_formatadas.append(data_formatada)

    if len(datas_formatadas) == 1:
        datas_concatenadas = datas_formatadas[0]
    elif len(datas_formatadas) == 2:
        datas_concatenadas = ' e '.join(datas_formatadas)
    elif len(datas_formatadas) > 2:
        datas_concatenadas = ', '.join(datas_formatadas[:-1]) + ' e ' + datas_formatadas[-1]
    else:
        datas_concatenadas = ''

    documento_html = documento_html.replace('{{datas_leilao}}',
                                            datas_concatenadas if datas_concatenadas else 'Sem datas definidas')

    modalidade = processo.leilao.modalidade if processo.leilao else 'Modalidade não encontrada'

    modalidade_texto = dict(Leilao.MODALIDADE_CHOICES).get(modalidade, modalidade)

    documento_html = documento_html.replace('{{modalidade}}', modalidade_texto)
    documento_html = documento_html.replace('{{modalidade_maiusculo}}', modalidade_texto.upper())

    datas_leilao = processo.leilao.datas_leilao.all() if processo.leilao else []
    local_tz = pytz.timezone('America/Sao_Paulo')

    for index, dataleilao in enumerate(datas_leilao, start=1):
        data_hora_local = dataleilao.data_hora.astimezone(local_tz)
        data_formatada = data_hora_local.strftime('%d/%m/%Y')
        hora_formatada = data_hora_local.strftime('%H:%M')

        etapa_texto = dict(DataLeilao.ETAPA_CHOICES).get(dataleilao.etapa, dataleilao.etapa)
        complemento_texto = dict(DataLeilao.COMPLEMENTO_CHOICES).get(dataleilao.complemento, dataleilao.complemento)

        if dataleilao.data_abertura:
            data_abertura_local = dataleilao.data_abertura.astimezone(local_tz)
            data_abertura_formatada = data_abertura_local.strftime('%d/%m/%Y')
            hora_abertura_formatada = data_abertura_local.strftime('%H:%M')

            frase_data = (
                f"{data_abertura_formatada}, às {hora_abertura_formatada} horas e se encerrará "
                f"dia {data_formatada}, às {hora_formatada} horas"
            )
        else:
            frase_data = f"{data_formatada} {complemento_texto} {hora_formatada} horas"

        placeholder = f"{{{{data_leilao_{index}}}}}"
        documento_html = documento_html.replace(placeholder, frase_data)

    if not datas_leilao:
        documento_html = documento_html.replace('{{data_leilao_1}}', 'Nenhuma data de leilão encontrada')

    if processo.venda_direta:
        vendadireta_texto = FrasePadrao.objects.filter(chave='vendadireta_texto').first()
        documento_html = documento_html.replace('{{vendadireta_texto}}',
                                                vendadireta_texto.descricao if vendadireta_texto else '')
    else:
        documento_html = documento_html.replace('{{vendadireta_texto}}', '')

    bem_imovel = processo.bens.filter(tipo_bem='Imoveis').first()
    if bem_imovel:
        imovel_condicao_venda = FrasePadrao.objects.filter(chave='imovel_condicao_venda').first()
        documento_html = documento_html.replace('{{imovel_condicao_venda}}',
                                                imovel_condicao_venda.descricao if imovel_condicao_venda else '')

        onus_imoveis = FrasePadrao.objects.filter(chave='onus_imoveis').first()
        documento_html = documento_html.replace('{{onus_imoveis}}', onus_imoveis.descricao if onus_imoveis else '')

        onus_imoveis_peticao = FrasePadrao.objects.filter(chave='onus_imoveis_peticao_juntada').first()
        documento_html = documento_html.replace('{{onus_imoveis_peticao_juntada}}',
                                                onus_imoveis_peticao.descricao if onus_imoveis_peticao else '')
    else:
        documento_html = documento_html.replace('{{imovel_condicao_venda}}', '')
        documento_html = documento_html.replace('{{onus_imoveis}}', '')
        documento_html = documento_html.replace('{{onus_imoveis_peticao_juntada}}', '')

    bem_automovel = processo.bens.filter(tipo_bem='Automoveis').first()
    if bem_automovel:
        onus_veiculos = FrasePadrao.objects.filter(chave='onus_veiculos').first()
        documento_html = documento_html.replace('{{onus_veiculos}}', onus_veiculos.descricao if onus_veiculos else '')

        onus_veiculos_peticao = FrasePadrao.objects.filter(chave='onus_veiculos_peticao_juntada').first()
        documento_html = documento_html.replace('{{onus_veiculos_peticao_juntada}}',
                                                onus_veiculos_peticao.descricao if onus_veiculos_peticao else '')
    else:
        documento_html = documento_html.replace('{{onus_veiculos}}', '')
        documento_html = documento_html.replace('{{onus_veiculos_peticao_juntada}}', '')

    # Remover todos os parágrafos vazios (<p>...</p>) que não contenham conteúdo
    #documento_html = re.sub(r'<p[^>]*>\s*(<br\s*/?>)?\s*</p>', '', documento_html)

    template_path = None

    if modelo.timbrado == 'vara' and processo.vara:
        vara = get_object_or_404(Vara, id=processo.vara_id)
        template_path = os.path.join(settings.MEDIA_ROOT, vara.timbrado.name)
        logger.debug(f"Template path for vara: {template_path}")
    elif modelo.timbrado == 'leiloeiro':
        leiloeiro = modelo.leiloeiro or processo.leiloeiro
        if leiloeiro:
            leiloeiro = get_object_or_404(Leiloeiro, id=leiloeiro.id)
            template_path = os.path.join(settings.MEDIA_ROOT, leiloeiro.timbrado.name)
            logger.debug(f"Template path for leiloeiro: {template_path}")
        else:
            logger.error("Leiloeiro não encontrado para o timbrado do modelo.")
            return HttpResponse("Erro: Leiloeiro não encontrado para o timbrado do modelo.", status=400)
    elif modelo.timbrado == 'justica' and modelo.justica:
        justica = get_object_or_404(Justica, id=modelo.justica_id)
        template_path = os.path.join(settings.MEDIA_ROOT, justica.timbrado.name)
        logger.debug(f"Template path for justica: {template_path}")
    else:
        template_path = os.path.join(settings.STATIC_ROOT, 'modelos', 'template_default.docx')
        logger.debug(f"Using default template path: {template_path}")

    if not os.path.exists(template_path):
        logger.error(f"Erro ao abrir o template: Package not found at '{template_path}'")
        return HttpResponse(f"Erro ao abrir o template: Package not found at '{template_path}'", status=500)

    try:
        doc = Document(template_path)
    except Exception as e:
        logger.error(f"Erro ao abrir o template: {e}")
        return HttpResponse(f"Erro ao abrir o template: {e}", status=500)

    destinatarios = processo.executados.all()

    if destinatarios:
        for i, paragraph in enumerate(doc.paragraphs):
            if '{{tabela_destinatarios}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{tabela_destinatarios}}', '')

                tabela = doc.add_table(rows=1, cols=2)
                tabela.style = 'Table Grid'
                tabela.alignment = WD_TABLE_ALIGNMENT.CENTER

                tbl = tabela._tbl
                tbl_pr = tbl.tblPr

                if tbl_pr is None:
                    tbl_pr = OxmlElement('w:tblPr')
                    tbl.append(tbl_pr)

                tbl_borders = OxmlElement('w:tblBorders')

                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border_element = OxmlElement(f'w:{border_name}')
                    border_element.set(qn('w:val'), 'single')
                    border_element.set(qn('w:sz'), '12')
                    border_element.set(qn('w:color'), '000000')
                    tbl_borders.append(border_element)

                tbl_pr.append(tbl_borders)

                hdr_cells = tabela.rows[0].cells
                hdr_cells[0].text = 'DESTINATÁRIO'
                hdr_cells[1].text = 'QUALIFICAÇÃO NOS AUTOS'

                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        run = paragraph.runs[0]
                        run.font.bold = True
                        run.font.size = Pt(12)

                for destinatario in destinatarios:
                    row_cells = tabela.add_row().cells
                    row_cells[0].text = destinatario.nome
                    row_cells[1].text = "Executado"

                    for cell in row_cells:
                        for paragraph in cell.paragraphs:
                            run = paragraph.runs[0]
                            run.font.size = Pt(12)

                paragraph._element.addnext(tabela._element)

                break
    else:
        for paragraph in doc.paragraphs:
            if '{{tabela_destinatarios}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{tabela_destinatarios}}', 'Sem destinatários registrados.')

    endereco_dinamico = processo.vara.endereco_completo() if processo.vara else 'Endereço não disponível'
    endereco_dinamico_maiusculo = endereco_dinamico.upper()
    nome_vara_maiusculo = processo.vara.nome.upper() if processo.vara else 'SEM VARA'
    cidade_vara_maiusculo = processo.vara.comarca.nome.upper() if processo.vara and processo.vara.comarca else 'SEM CIDADE'
    telefone_vara = processo.vara.telefone if processo.vara and processo.vara.telefone else 'Telefone não disponível'
    email_vara = processo.vara.email if processo.vara and processo.vara.email else 'E-mail não disponível'

    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            replace_placeholder_in_paragraph(paragraph, '{{endereco_completo_vara}}', endereco_dinamico)
            replace_placeholder_in_paragraph(paragraph, '{{endereco_completo_vara_maiusculo}}',
                                             endereco_dinamico_maiusculo, bold=True)
            replace_placeholder_in_paragraph(paragraph, '{{nome_vara_maiusculas}}', nome_vara_maiusculo, bold=True)
            replace_placeholder_in_paragraph(paragraph, '{{cidade_vara_maiusculo}}', cidade_vara_maiusculo, bold=True)
            replace_placeholder_in_paragraph(paragraph, '{{telefone_vara}}', telefone_vara)
            replace_placeholder_in_paragraph(paragraph, '{{email_vara}}', email_vara)
    soup = BeautifulSoup(documento_html, 'html.parser')
    for elem in soup.find_all(['p', 'table']):
        if elem.name == 'p':
            paragraph = doc.add_paragraph()
            alignment = elem.get('style', '')
            if 'text-align: center' in alignment:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif 'text-align: right' in alignment:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif 'text-align: justify' in alignment:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            if 'text-indent' in alignment:
                indent_match = re.search(r'text-indent:\s*(\d+)px;?', alignment)
                if indent_match:
                    indent_value_px = int(indent_match.group(1))
                    indent_value_pt = indent_value_px * 0.75
                    indent_value_pt += 40
                    paragraph.paragraph_format.first_line_indent = Pt(indent_value_pt)

            for content in elem.contents:
                if isinstance(content, Tag):
                    if content.name == 'br':
                        paragraph.add_run().add_break()
                    elif content.name == 'strong':
                        run = paragraph.add_run(content.get_text())
                        run.bold = True
                    elif content.name == 'em':
                        run = paragraph.add_run(content.get_text())
                        run.italic = True
                    else:
                        run = paragraph.add_run(content.get_text())
                elif isinstance(content, NavigableString):
                    run = paragraph.add_run(str(content))

        elif elem.name == 'table':
            gerar_tabela_dinamica(doc, elem)

        penhoras_externas = processo.penhoras_externas.all()

        if penhoras_externas:
            for paragraph in doc.paragraphs:
                if '{{tabela_penhoras_externas}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{tabela_penhoras_externas}}', '')

                    tabela = doc.add_table(rows=1, cols=3)
                    tabela.allow_autofit = False

                    tbl = tabela._tbl
                    tbl_pr = tbl.tblPr

                    if tbl_pr is None:
                        tbl_pr = OxmlElement('w:tblPr')
                        tbl.append(tbl_pr)

                    tbl_borders = OxmlElement('w:tblBorders')

                    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                        border_element = OxmlElement(f'w:{border_name}')
                        border_element.set(qn('w:val'), 'single')
                        border_element.set(qn('w:sz'), '12')
                        border_element.set(qn('w:color'), '000000')
                        tbl_borders.append(border_element)

                    tbl_pr.append(tbl_borders)

                    hdr_cells = tabela.rows[0].cells
                    hdr_cells[0].text = 'NÚMERO'
                    hdr_cells[1].text = 'VARA/COMARCA'
                    hdr_cells[2].text = 'EXEQUENTE'

                    for cell in hdr_cells:
                        for paragraph in cell.paragraphs:
                            run = paragraph.runs[0]
                            run.font.bold = True
                            run.font.size = Pt(12)

                    for penhora in penhoras_externas:
                        row_cells = tabela.add_row().cells
                        row_cells[0].text = penhora.numero if penhora.numero else 'Número não especificado'
                        row_cells[1].text = penhora.vara.nome if penhora.vara else 'Sem Vara'
                        row_cells[2].text = penhora.exequente.nome if penhora.exequente else 'Sem Exequente'

                        for cell in row_cells:
                            for paragraph in cell.paragraphs:
                                run = paragraph.runs[0]
                                run.font.size = Pt(12)

                    break
        else:
            for paragraph in doc.paragraphs:
                if '{{tabela_penhoras_externas}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{tabela_penhoras_externas}}',
                                                            'Sem penhoras externas registradas.')

    destinatarios = processo.executados.all()

    if destinatarios:
        for i, paragraph in enumerate(doc.paragraphs):
            if '{{tabela_destinatarios}}' in paragraph.text:
                paragraph.clear()

                tabela = doc.add_table(rows=1, cols=2)
                tabela.allow_autofit = False

                tbl = tabela._tbl
                tbl_pr = tbl.tblPr

                if tbl_pr is None:
                    tbl_pr = OxmlElement('w:tblPr')
                    tbl.append(tbl_pr)

                tbl_borders = OxmlElement('w:tblBorders')

                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border_element = OxmlElement(f'w:{border_name}')
                    border_element.set(qn('w:val'), 'single')
                    border_element.set(qn('w:sz'), '12')
                    border_element.set(qn('w:color'), '000000')
                    tbl_borders.append(border_element)

                tbl_pr.append(tbl_borders)

                hdr_cells = tabela.rows[0].cells
                hdr_cells[0].text = 'DESTINATÁRIO'
                hdr_cells[1].text = 'QUALIFICAÇÃO NOS AUTOS'

                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        run = paragraph.runs[0]
                        run.font.bold = True
                        run.font.size = Pt(12)

                for destinatario in destinatarios:
                    row_cells = tabela.add_row().cells
                    row_cells[0].text = destinatario.nome
                    row_cells[1].text = 'Executado'

                    for cell in row_cells:
                        for paragraph in cell.paragraphs:
                            run = paragraph.runs[0]
                            run.font.size = Pt(12)

                doc.paragraphs[i]._element.addnext(tabela._element)

                break
    else:
        for paragraph in doc.paragraphs:
            if '{{tabela_destinatarios}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{tabela_destinatarios}}', 'Sem destinatários registrados.')

    output_filename = os.path.join(settings.MEDIA_ROOT, f"documento_{modelo.nome_modelo}_{processo.numero}.docx")
    doc.save(output_filename)

    with open(output_filename, 'rb') as docx_file:
        response = HttpResponse(
            docx_file.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{os.path.basename(output_filename)}"'
    return response


class ProcessoViewSet(viewsets.ModelViewSet):
    """Exibindo todos processos"""
    queryset = Processo.objects.all()
    serializer_class = ProcessoSerializer


def editar_documento(request, processo_id):
    processo = get_object_or_404(Processo, pk=processo_id)

    modelo = MyModel.objects.filter(leiloeiro=processo.leiloeiro, tipo_arquivo='sugestao_datas').first()

    if modelo is None:
        return HttpResponse("Modelo de documento não encontrado.", status=404)

    documento = modelo.content
    documento = documento.replace('{{numero_processo}}', str(processo.numero))

    return render(request, 'editar_documento.html', {'documento': documento})


def get_modelo_documento(tipo_documento_nome, vara=None, leiloeiro=None):
    tipo_documento = TipoDocumento.objects.get(nome=tipo_documento_nome)
    modelo = MyModel.objects.filter(tipo_documento=tipo_documento, vara=vara).first()
    if not modelo:
        modelo = MyModel.objects.filter(tipo_documento=tipo_documento, leiloeiro=leiloeiro).first()
    if not modelo:
        modelo = MyModel.objects.filter(tipo_documento=tipo_documento, e_modelo_padrao=True).first()
    return modelo


def get_processo_details(request, card_id):
    try:
        card = Card.objects.select_related(
            'processo__vara',
            'processo__leiloeiro',
            'processo__responsavel',
            'processo__leilao'
        ).prefetch_related(
            Prefetch('processo__leilao__datas_leilao')
        ).get(id=card_id)

        processo = card.processo
        leilao = processo.leilao

        datas = [data.data_hora.strftime("%d/%m/%Y %H:%M") for data in leilao.datas_leilao.all()] if leilao else []

        data = {
            'numero': processo.numero if processo else "Sem Número",
            'vara': processo.vara.nome if processo and processo.vara else "Sem Vara",
            'leilao': leilao.nome if leilao else "Sem Leilão",
            'datas': datas,
            'leiloeiro': processo.leiloeiro.nome if processo and processo.leiloeiro else "Sem Leiloeiro",
            'responsavel': processo.responsavel.get_full_name() if processo and processo.responsavel else "Sem Responsável",
        }
        return JsonResponse(data)
    except Card.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'Card não encontrado'}, status=404)


@login_required
@require_POST
def add_comment(request, card_id):
    card = get_object_or_404(Card, id=card_id)
    content = request.POST.get('content')
    mentioned_usernames = [word[1:] for word in content.split() if word.startswith('@')]

    if content:
        comment = Comment.objects.create(card=card, author=request.user, content=content)

        if mentioned_usernames:
            mentioned_users = User.objects.filter(username__in=mentioned_usernames)
            for user in mentioned_users:
                notify.send(request.user, recipient=user, verb='mentioned you in a comment', target=card, description=content)

        data = {
            'author': comment.author.get_full_name(),
            'content': comment.content,
            'created_at': comment.created_at.strftime('%d/%m/%Y %H:%M')
        }
        return JsonResponse(data)

    return JsonResponse({'error': 'Content cannot be empty'}, status=400)


class UserAutocomplete(View):
    def get(self, request, *args, **kwargs):
        if not request.user.is_authenticated:
            return JsonResponse({'results': []})

        term = request.GET.get('term', '')
        users = User.objects.filter(username__icontains=term).values('username')
        results = list(users)
        return JsonResponse({'results': results})


def get_modelos(request, processo_id):
    processo = get_object_or_404(Processo, pk=processo_id)

    # Obtenha o ID da lista do parâmetro da URL
    list_id = request.GET.get('list_id')
    if not list_id:
        return JsonResponse({'modelos': []})

    # Mapear tipos de documentos para listas
    documento_lista_map = {
        1: [1],  # Documento de ID 1 deve aparecer apenas na Lista de ID 1
        2: [4],
        3: [4],
    }

    tipos_permitidos_ids = [
        tipo_id for tipo_id, lista_ids in documento_lista_map.items() if int(list_id) in lista_ids
    ]

    if not tipos_permitidos_ids:
        return JsonResponse({'modelos': []})

    modelos_vara = MyModel.objects.filter(tipo_documento__id__in=tipos_permitidos_ids, vara=processo.vara)
    modelos_leiloeiro = MyModel.objects.filter(tipo_documento__id__in=tipos_permitidos_ids, leiloeiro=processo.leiloeiro)

    modelos = (modelos_vara | modelos_leiloeiro).distinct()

    modelos = modelos.filter(leiloeiro__isnull=True) | modelos.filter(leiloeiro=processo.leiloeiro)

    modelos_data = [
        {
            'id': modelo.id,
            'nome_modelo': modelo.nome_modelo,
            'tipo_documento': {'nome': modelo.tipo_documento.nome}
        }
        for modelo in modelos
    ]

    return JsonResponse({'modelos': modelos_data})

@receiver(post_save, sender=Card)
def atribuir_grupos_automaticamente(sender, instance, created, **kwargs):
    if not created:
        estado_vara = instance.processo.vara.comarca.estado
        grupos = GrupoResponsavel.objects.filter(estados=estado_vara)
        current_grupos = set(instance.grupos_responsaveis.all())
        new_grupos = set(grupos)
        if current_grupos != new_grupos:
            instance.grupos_responsaveis.set(grupos)
            instance.save()


class GruposResponsaveisFilter(SimpleListFilter):
    title = 'Grupos Responsáveis'
    parameter_name = 'grupo'

    def lookups(self, request, model_admin):
        return [('todos', 'Todos os Grupos')] + [(grupo.id, grupo.nome) for grupo in GrupoResponsavel.objects.all()]

    def queryset(self, request, queryset):
        if self.value() == 'todos':
            return queryset
        elif self.value():
            return queryset.filter(grupos_responsaveis__id=self.value())
        return queryset


def get_arquivos(request, processo_id):
    arquivos = ProcessoArquivo.objects.filter(processo_id=processo_id)
    arquivos_data = [
        {
            'id': arquivo.id,
            'nome': arquivo.nome_arquivo,
            'url': arquivo.arquivo.url
        }
        for arquivo in arquivos
    ]
    return JsonResponse({'arquivos': arquivos_data})

def render_template(template_content, context):
    """
    Substitui os placeholders no template pelo contexto fornecido.
    """
    for key, value in context.items():
        placeholder = f'{{{{{key}}}}}'
        template_content = template_content.replace(placeholder, str(value))
    return template_content

def send_email_from_template(template_id, processo_id):
    """
    Envia um email baseado em um template armazenado no banco de dados.
    """
    template = get_object_or_404(EmailTemplate, pk=template_id)
    processo = get_object_or_404(Processo, pk=processo_id)

    card = Card.objects.filter(processo=processo).first()

    if not card:
        raise ValueError("Card relacionado ao processo não foi encontrado.")

    context = {
        'numero_processo': processo.numero,
        'nome_exequente': ', '.join([exequente.nome for exequente in processo.exequentes.all()]),
        'nome_executado': ', '.join([executado.nome for executado in processo.executados.all()]),
        'nome_vara': processo.vara.nome if processo.vara else 'Sem Vara',
        'nome_vara_maiusculas': processo.vara.nome.upper() if processo.vara else 'SEM VARA',
        'nome_leiloeiro': processo.leiloeiro.nome if processo.leiloeiro else 'Sem Leiloeiro',
        'nome_leiloeiro_maiusculo': processo.leiloeiro.nome.upper() if processo.leiloeiro else 'SEM LEILOEIRO',
        'classe': processo.classe_processual.nome if processo.classe_processual else 'Sem Classe Processual',
        'cidade_vara': processo.vara.comarca.nome if processo.vara and processo.vara.comarca else 'Sem Cidade',
        'cidade_vara_maiusculo': processo.vara.comarca.nome.upper() if processo.vara and processo.vara.comarca else 'SEM CIDADE',
        'justica': processo.vara.justica.nome if processo.vara and processo.vara.justica else 'Sem Justiça',
        'justica_maiusculo': processo.vara.justica.nome.upper() if processo.vara and processo.vara.justica else 'SEM JUSTIÇA',
        'uf': processo.vara.comarca.estado.uf if processo.vara and processo.vara.comarca and processo.vara.comarca.estado else 'UF não definida',
        'site': processo.leiloeiro.site if processo.leiloeiro else 'Sem Site',
        'data_hoje': timezone.now().strftime('%d/%m/%Y'),
        'card_id': card.id if card else '',
        'board_id': card.list.board.id if card and card.list and card.list.board else '',
        'site_url': settings.SITE_URL,
    }

    # Processar datas de leilão
    datas_leilao = processo.leilao.datas_leilao.all() if processo.leilao else []
    datas_leilao_texto = ""
    local_tz = pytz.timezone('America/Sao_Paulo')
    for dataleilao in datas_leilao:
        data_hora_local = dataleilao.data_hora.astimezone(local_tz)
        data_formatada = data_hora_local.strftime('%d/%m/%Y')
        hora_formatada = data_hora_local.strftime('%H:%M')
        etapa_texto = dict(DataLeilao.ETAPA_CHOICES).get(dataleilao.etapa, dataleilao.etapa)
        tipo_data_texto = dict(DataLeilao.TIPO_CHOICES).get(dataleilao.tipo_data, dataleilao.tipo_data)
        complemento_texto = dict(DataLeilao.COMPLEMENTO_CHOICES).get(dataleilao.complemento, dataleilao.complemento)
        datas_leilao_texto += f"{etapa_texto} - {tipo_data_texto}: {data_formatada}, {complemento_texto} {hora_formatada} horas<br>"

    context['datas_leilao'] = datas_leilao_texto.rstrip(
        '<br>') if datas_leilao_texto else 'Nenhuma data de leilão encontrada'

    leiloeiro = processo.leiloeiro
    if leiloeiro:
        matricula = leiloeiro.matriculas.filter(
            estado=processo.vara.comarca.estado).first() if processo.vara and processo.vara.comarca else None
        if matricula:
            context['numero_matricula'] = matricula.matricula
            context['cidade_leiloeiro'] = matricula.cidade.nome if matricula.cidade else 'Cidade não cadastrada'
        else:
            context['numero_matricula'] = 'Matrícula não encontrada'
            context['cidade_leiloeiro'] = 'Matrícula não encontrada'
    else:
        context['numero_matricula'] = 'Leiloeiro não encontrado'
        context['cidade_leiloeiro'] = 'Leiloeiro não encontrado'

    email_html = render_template(template.content, context)

    status = 'failed'
    sent_at = None
    try:
        send_mail(
            template.subject,
            '',
            'noreply@leiloesjudiciais.com.br',
            [recipient.email for recipient in template.recipients.all()],
            fail_silently=False,
            html_message=email_html
        )
        status = 'sent'
        sent_at = timezone.now()
    except Exception as e:
        print(f"Failed to send email: {e}")

    EmailLog.objects.create(
        subject=template.subject,
        message=email_html,
        recipient=', '.join([r.email for r in template.recipients.all()]),
        status=status,
        sent_at=sent_at
    )


@csrf_exempt
@require_POST
def move_card(request):
    data = json.loads(request.body)
    card_id = data.get('card_id')
    new_list_id = data.get('new_list_id')

    try:
        card = Card.objects.get(id=card_id)
        new_list = List.objects.get(id=new_list_id)

        if card.list_id != new_list_id:
            card.list = new_list
            card.save()

            email_templates = new_list.email_templates.all()
            for template in email_templates:
                try:
                    send_email_from_template(template.id, card.processo.id)
                except Exception as e:
                    print(f"Failed to send email: {e}")

        return JsonResponse({'status': 'success'})

    except Card.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'Card not found'}, status=404)
    except List.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'List not found'}, status=404)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@require_POST
def add_checklist_item(request, card_id):
    card = get_object_or_404(Card, id=card_id)
    item_name = request.POST.get('name')
    if item_name:
        ChecklistItem.objects.create(card=card, name=item_name)
        return JsonResponse({'status': 'success'})
    return JsonResponse({'status': 'error', 'message': 'Nome da tarefa é obrigatório'}, status=400)

@require_POST
def toggle_checklist_item(request, item_id):
    item = get_object_or_404(ChecklistItem, id=item_id)
    item.completed = not item.completed
    item.save()
    return JsonResponse({'status': 'success'})

@require_POST
def remove_checklist_item(request, item_id):
    item = get_object_or_404(ChecklistItem, id=item_id)
    item.delete()
    return JsonResponse({'status': 'success'})


def get_checklist(request, card_id):
    tasks = ChecklistItem.objects.filter(card_id=card_id).values('id', 'name', 'completed')
    tasks_list = list(tasks)

    return JsonResponse({'tasks': tasks_list})


def get_modelos_kanban(request, processo_id):
    processo = get_object_or_404(Processo, pk=processo_id)
    vara = processo.vara
    justica = vara.justica if vara else None

    modelos = MyModel.objects.filter(
        Q(vara__isnull=True) | Q(vara=vara),
        Q(justica__isnull=True) | Q(justica=justica),
        Q(leiloeiro__isnull=True) | Q(leiloeiro=processo.leiloeiro)
    )

    modelos_data = [
        {
            'id': modelo.id,
            'nome_modelo': modelo.nome_modelo,
            'tipo_documento': {
                'nome': modelo.tipo_documento.nome
            }
        } for modelo in modelos
    ]
    return JsonResponse({'modelos': modelos_data})

class LeilaoAutocomplete(autocomplete.Select2QuerySetView):
    def get_queryset(self):
        if not self.request.user.is_authenticated:
            return Leilao.objects.none()

        queryset = Leilao.objects.filter(datas_leilao__data_hora__gte=now()).distinct()

        if self.q:
            queryset = queryset.filter(nome__icontains=self.q)

        return queryset


def get_filtered_agregados_processos(request):
    estado = request.GET.get('estado')
    leiloeiro = request.GET.get('leiloeiro')
    dias_antes_leilao = request.GET.get('dias_antes_leilao')
    dias_depois_cadastro = request.GET.get('dias_depois_cadastro')
    intervalo_inicio = request.GET.get('intervalo_inicio')
    intervalo_fim = request.GET.get('intervalo_fim')
    coluna_kanban = request.GET.get('coluna_kanban')

    cards_agregados = CardsAgregados.objects.filter(card__list_id=coluna_kanban) if coluna_kanban else CardsAgregados.objects.all()

    if dias_antes_leilao:
        cards_agregados = cards_agregados.filter(
            card__leilao__datas_leilao__data_hora__lte=timezone.now() + timedelta(days=int(dias_antes_leilao))
        )

    processos_ids = cards_agregados.values_list('processo_id', flat=True)

    processos = Processo.objects.filter(id__in=processos_ids)

    if estado:
        processos = processos.filter(vara__cidade__estado__nome=estado)
    if leiloeiro:
        processos = processos.filter(leiloeiro__nome=leiloeiro)
    if dias_depois_cadastro:
        processos = processos.filter(data_cadastro__gte=timezone.now() - timedelta(days=int(dias_depois_cadastro)))
    if intervalo_inicio and intervalo_fim:
        processos = processos.filter(data_cadastro__range=[intervalo_inicio, intervalo_fim])

    return processos.distinct()


def get_filtered_cards_board_3(request):
    estado = request.GET.get('estado')
    leiloeiro = request.GET.get('leiloeiro')
    dias_antes_leilao = request.GET.get('dias_antes_leilao')
    intervalo_inicio = request.GET.get('intervalo_inicio')
    intervalo_fim = request.GET.get('intervalo_fim')
    coluna_kanban = request.GET.get('coluna_kanban')

    cards = Card.objects.filter(list__board_id=3).distinct()

    if estado:
        cards = cards.filter(vara__comarca__estado__nome=estado)

    if leiloeiro:
        cards = cards.filter(leiloeiro__nome=leiloeiro)

    if dias_antes_leilao:
        dias_antes = timezone.now() + timedelta(days=int(dias_antes_leilao))
        cards = cards.filter(leilao__datas_leilao__data_hora__lte=dias_antes)

    if intervalo_inicio and intervalo_fim:
        cards = cards.filter(
            leilao__datas_leilao__data_hora__range=[intervalo_inicio, intervalo_fim]
        )

    if coluna_kanban:
        cards = cards.filter(list_id=coluna_kanban)

    return cards


def download_pdf(request):
    try:
        board_id = request.GET.get('board_id')

        if not board_id:
            return HttpResponse("Parâmetro board_id ausente ou inválido.", status=400)

        if board_id == '3':
            estado = request.GET.get('estado')
            leiloeiro = request.GET.get('leiloeiro')
            dias_antes_leilao = request.GET.get('dias_antes_leilao')
            intervalo_inicio = request.GET.get('intervalo_inicio')
            intervalo_fim = request.GET.get('intervalo_fim')
            coluna_kanban = request.GET.get('coluna_kanban')

            cards = Card.objects.filter(list__board_id=3)

            if estado:
                cards = cards.filter(vara__comarca__estado__nome=estado)
            if leiloeiro:
                cards = cards.filter(leiloeiro__nome=leiloeiro)
            if dias_antes_leilao:
                dias_antes = timezone.now() + timedelta(days=int(dias_antes_leilao))
                cards = cards.filter(leilao__datas_leilao__data_hora__lte=dias_antes)
            if intervalo_inicio and intervalo_fim:
                cards = cards.filter(leilao__datas_leilao__data_hora__range=[intervalo_inicio, intervalo_fim])
            if coluna_kanban:
                cards = cards.filter(list_id=coluna_kanban)

            for card in cards:
                datas_leilao = card.leilao.datas_leilao.all().order_by('data_hora') if card.leilao else None
                if datas_leilao:
                    card.primeira_data = datas_leilao.first().data_hora.strftime('%d/%m/%Y %H:%M')
                    card.ultima_data = datas_leilao.last().data_hora.strftime('%d/%m/%Y %H:%M')
                else:
                    card.primeira_data = 'N/A'
                    card.ultima_data = 'N/A'

            html_string = render_to_string('extracao_pdf/pdf_template_board_3.html', {'cards': cards})

        else:
            processos = get_filtered_processos(request)

            for processo in processos:
                datas_leilao = processo.leilao.datas_leilao.all().order_by('data_hora')
                if datas_leilao.exists():
                    processo.primeira_data = make_naive(datas_leilao.first().data_hora).strftime('%d/%m/%Y %H:%M')
                    processo.ultima_data = make_naive(datas_leilao.last().data_hora).strftime('%d/%m/%Y %H:%M')
                else:
                    processo.primeira_data = 'N/A'
                    processo.ultima_data = 'N/A'

            html_string = render_to_string('extracao_pdf/pdf_template.html', {'processos': processos})

        pdf_file = BytesIO()
        if board_id == '3':
            HTML(string=html_string).write_pdf(pdf_file, stylesheets=[CSS(string="@page { size: A4 landscape; }")])
        else:
            HTML(string=html_string).write_pdf(pdf_file)

        pdf_file.seek(0)

        response = HttpResponse(pdf_file, content_type='application/pdf')
        response['Content-Disposition'] = f'inline; filename="relatorio_board_{board_id}.pdf"'
        return response

    except Exception as e:
        print(f"Erro ao gerar PDF: {e}")

        return HttpResponse("Erro ao gerar o PDF.", status=500)



import pandas as pd
from io import BytesIO
from django.http import HttpResponse
from django.utils.timezone import make_naive
from datetime import timedelta

def download_excel(request):
    try:
        board_id = request.GET.get('board_id')

        if board_id == '3':
            cards = get_filtered_agregados_processos(request)

            data = []
            for card in cards.distinct():
                datas_leilao = card.leilao.datas_leilao.all().order_by('data_hora')
                primeira_data = datas_leilao.first().data_hora.strftime('%d/%m/%Y %H:%M') if datas_leilao.exists() else 'N/A'
                ultima_data = datas_leilao.last().data_hora.strftime('%d/%m/%Y %H:%M') if datas_leilao.exists() else 'N/A'

                data.append({
                    'ID do Card': card.id,
                    'Título do Leilão': card.leilao.nome if card.leilao else 'N/A',
                    'Vara': card.vara.nome if card.vara else 'N/A',
                    'Leiloeiro': card.leiloeiro.nome if card.leiloeiro else 'N/A',
                    'Primeira Data do Leilão': primeira_data,
                    'Última Data do Leilão': ultima_data,
                })

            df = pd.DataFrame(data)
            sheet_name = 'Relatorio Board 3'

        else:
            processos = get_filtered_processos(request)

            data = []
            for processo in processos:
                datas_leilao = processo.leilao.datas_leilao.all().order_by('data_hora')
                if datas_leilao.exists():
                    primeira_data = make_naive(datas_leilao.first().data_hora).strftime('%d/%m/%Y %H:%M')
                    ultima_data = make_naive(datas_leilao.last().data_hora).strftime('%d/%m/%Y %H:%M')
                else:
                    primeira_data = 'N/A'
                    ultima_data = 'N/A'

                data.append({
                    'ID': processo.id,
                    'Número': processo.numero,
                    'Vara': processo.vara.nome if processo.vara else 'N/A',
                    'Leiloeiro': processo.leiloeiro.nome if processo.leiloeiro else 'N/A',
                    'Data de Cadastro': processo.data_cadastro.strftime('%d/%m/%Y') if processo.data_cadastro else 'N/A',
                    'Primeira Data do Leilão': primeira_data,
                    'Última Data do Leilão': ultima_data
                })

            df = pd.DataFrame(data)
            sheet_name = 'Processos'

        excel_file = BytesIO()
        with pd.ExcelWriter(excel_file, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name=sheet_name)
            worksheet = writer.sheets[sheet_name]

            for column_cells in worksheet.columns:
                max_length = max(len(str(cell.value)) for cell in column_cells)
                adjusted_width = (max_length + 2)
                worksheet.column_dimensions[column_cells[0].column_letter].width = adjusted_width

        excel_file.seek(0)

        response = HttpResponse(excel_file, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename=relatorio_board_{board_id}.xlsx'
        return response

    except Exception as e:
        print(f"Erro ao gerar Excel: {e}")
        return HttpResponse("Erro ao gerar o Excel.", status=500)


def filtrar_relatorios(request):
    processos = get_filtered_processos(request)

    processos_listados = []
    for processo in processos:
        card = Card.objects.filter(processo_id=processo.id).first()
        coluna_kanban = card.list.name if card and card.list else 'Sem Coluna'

        processos_listados.append({
            'id': processo.id,
            'numero': processo.numero,
            'vara__nome': processo.vara.nome if processo.vara else 'N/A',
            'leilao__nome': processo.leilao.nome if processo.leilao else 'N/A',
            'leiloeiro__nome': processo.leiloeiro.nome if processo.leiloeiro else 'N/A',
            'data_cadastro': processo.data_cadastro.strftime('%d/%m/%Y') if processo.data_cadastro else 'N/A',
            'primeira_data': processo.primeira_data,
            'ultima_data': processo.ultima_data,
            'coluna_kanban': coluna_kanban
        })

    return JsonResponse({'processos_listados': processos_listados})


def get_users(request):
    if not request.user.is_authenticated:
        return JsonResponse({'results': []}, status=401)

    term = request.GET.get('term', '')
    users = User.objects.filter(first_name__icontains=term)
    results = list(users.values('id', 'first_name', 'last_name'))

    return JsonResponse({'results': results})

@csrf_exempt
def update_responsavel(request, card_id):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            novo_responsavel_id = data.get('responsavel_id')
            card = Card.objects.get(id=card_id)
            card.responsavel_id = novo_responsavel_id
            card.save()
            return JsonResponse({'status': 'success'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)})
    return JsonResponse({'status': 'error', 'message': 'Método inválido'})

class ComarcaAutocomplete(autocomplete.Select2QuerySetView):
    def get_queryset(self):
        if not self.request.user.is_authenticated:
            return Cidade.objects.none()

        qs = Cidade.objects.filter(comarca=True)

        if self.q:
            qs = qs.filter(nome__icontains=self.q)

        return qs

def get_processos_agregados(request, card_id):
    processos = CardsAgregados.objects.filter(card_id=card_id).select_related('processo')
    data = {
        'processos': [
            {
                'id': processo.processo.id,
                'numero': processo.processo.numero
            }
            for processo in processos
        ]
    }
    return JsonResponse(data)

def admin_relatorios_board_1(request):
    processos = Processo.objects.none()

    estado = request.GET.get('estado')
    leiloeiro = request.GET.get('leiloeiro')
    dias_antes_leilao = request.GET.get('dias_antes_leilao')
    dias_depois_cadastro = request.GET.get('dias_depois_cadastro')
    intervalo_inicio = request.GET.get('intervalo_inicio')
    intervalo_fim = request.GET.get('intervalo_fim')
    coluna_kanban = request.GET.get('coluna_kanban')

    estados = Estado.objects.all()
    leiloeiros = Leiloeiro.objects.all()
    colunas_kanban = List.objects.filter(board_id=1)

    if estado or leiloeiro or dias_antes_leilao or dias_depois_cadastro or (intervalo_inicio and intervalo_fim) or coluna_kanban:
        processos = get_filtered_processos(request)

    context = {
        'processos': processos,
        'estados': estados,
        'leiloeiros': leiloeiros,
        'colunas_kanban': colunas_kanban,
    }

    return render(request, 'extracao_pdf/relatorios_board_1.html', context)

def admin_relatorios_board_3(request):
    estado = request.GET.get('estado')
    leiloeiro = request.GET.get('leiloeiro')
    dias_antes_leilao = request.GET.get('dias_antes_leilao')
    dias_depois_cadastro = request.GET.get('dias_depois_cadastro')
    intervalo_inicio = request.GET.get('intervalo_inicio')
    intervalo_fim = request.GET.get('intervalo_fim')
    coluna_kanban = request.GET.get('coluna_kanban')

    cards_agregados = CardsAgregados.objects.filter(card__list_id=coluna_kanban) if coluna_kanban else CardsAgregados.objects.all()

    processos_ids = cards_agregados.values_list('processo_id', flat=True)
    processos = Processo.objects.filter(id__in=processos_ids)

    if estado:
        processos = processos.filter(vara__cidade__estado__nome=estado)
    if leiloeiro:
        processos = processos.filter(leiloeiro__nome=leiloeiro)
    if dias_antes_leilao:
        processos = processos.filter(primeira_data__lte=timezone.now() + timedelta(days=int(dias_antes_leilao)))
    if dias_depois_cadastro:
        processos = processos.filter(data_cadastro__gte=timezone.now() - timedelta(days=int(dias_depois_cadastro)))
    if intervalo_inicio and intervalo_fim:
        processos = processos.filter(data_cadastro__range=[intervalo_inicio, intervalo_fim])

    context = {
        'processos': processos,
        'estados': Estado.objects.all(),
        'leiloeiros': Leiloeiro.objects.all(),
        'colunas_kanban': List.objects.filter(board_id=3)
    }
    return render(request, 'extracao_pdf/relatorios_board_3.html', context)


def filtrar_relatorios_board_3(request):
    estado = request.GET.get('estado')
    leiloeiro = request.GET.get('leiloeiro')
    dias_antes_leilao = request.GET.get('dias_antes_leilao')
    dias_depois_cadastro = request.GET.get('dias_depois_cadastro')
    intervalo_inicio = request.GET.get('intervalo_inicio')
    intervalo_fim = request.GET.get('intervalo_fim')
    coluna_kanban = request.GET.get('coluna_kanban')

    cards_agregados = CardsAgregados.objects.filter(
        card__list_id=coluna_kanban) if coluna_kanban else CardsAgregados.objects.all()

    processos_ids = cards_agregados.values_list('processo_id', flat=True)
    processos = Processo.objects.filter(id__in=processos_ids)

    if estado:
        processos = processos.filter(vara__cidade__estado__nome=estado)
    if leiloeiro:
        processos = processos.filter(leiloeiro__nome=leiloeiro)
    if dias_antes_leilao:
        processos = processos.filter(data_leilao__lte=timezone.now() + timedelta(days=int(dias_antes_leilao)))
    if dias_depois_cadastro:
        processos = processos.filter(data_cadastro__gte=timezone.now() - timedelta(days=int(dias_depois_cadastro)))
    if intervalo_inicio and intervalo_fim:
        processos = processos.filter(data_cadastro__range=[intervalo_inicio, intervalo_fim])

    processos_listados = []
    for processo in processos:
        card = processo.cardsagregados_set.first().card
        leilao = card.leilao if card else None

        primeira_data = 'N/A'
        ultima_data = 'N/A'
        if leilao:
            datas_leilao = DataLeilao.objects.filter(leilao=leilao).order_by('data_hora')
            if datas_leilao.exists():
                primeira_data = datas_leilao.first().data_hora.strftime('%d/%m/%Y %H:%M')
                ultima_data = datas_leilao.last().data_hora.strftime('%d/%m/%Y %H:%M')

        coluna_kanban_name = card.list.name if card and card.list else 'N/A'

        processos_listados.append({
            'id': processo.id,
            'numero': processo.numero,
            'vara__nome': processo.vara.nome if processo.vara else 'N/A',
            'leilao__nome': leilao.nome if leilao else 'N/A',
            'leiloeiro__nome': processo.leiloeiro.nome if processo.leiloeiro else 'N/A',
            'data_cadastro': processo.data_cadastro.strftime('%d/%m/%Y') if processo.data_cadastro else 'N/A',
            'primeira_data': primeira_data,
            'ultima_data': ultima_data,
            'coluna_kanban': coluna_kanban_name
        })

    return JsonResponse({'processos_listados': processos_listados})


def relatorio_simples_board_3(request):
    estado = request.GET.get('estado')
    leiloeiro = request.GET.get('leiloeiro')
    dias_antes_leilao = request.GET.get('dias_antes_leilao')
    intervalo_inicio = request.GET.get('intervalo_inicio')
    intervalo_fim = request.GET.get('intervalo_fim')
    coluna_kanban = request.GET.get('coluna_kanban')

    if any([estado, leiloeiro, dias_antes_leilao, intervalo_inicio, intervalo_fim, coluna_kanban]):
        cards = Card.objects.filter(list__board_id=3)

        cards = cards.annotate(primeira_data=Min('leilao__datas_leilao__data_hora'))

        if estado:
            cards = cards.filter(vara__comarca__estado__nome=estado)
        if leiloeiro:
            cards = cards.filter(leiloeiro__nome=leiloeiro)
        if dias_antes_leilao:
            dias_antes = timezone.now() + timedelta(days=int(dias_antes_leilao))
            cards = cards.filter(primeira_data__lte=dias_antes)
        if intervalo_inicio and intervalo_fim:
            cards = cards.filter(primeira_data__range=[intervalo_inicio, intervalo_fim])
        if coluna_kanban:
            cards = cards.filter(list_id=coluna_kanban)

        cards = cards.distinct()
    else:
        cards = Card.objects.none()

    estados = Estado.objects.all()
    leiloeiros = Leiloeiro.objects.all()
    colunas_kanban = List.objects.filter(board_id=3)

    return render(request, 'extracao_pdf/relatorio_simples_board_3.html', {
        'cards': cards,
        'estados': estados,
        'leiloeiros': leiloeiros,
        'colunas_kanban': colunas_kanban,
    })

@require_POST
def save_publicjud(request, card_id):
    card = get_object_or_404(Card, id=card_id)
    data = json.loads(request.body.decode('utf-8'))
    novo_identificador_publicjud = data.get('identificador_publicjud')

    if novo_identificador_publicjud:
        card.identificador_publicjud = novo_identificador_publicjud
        card.save()
        return JsonResponse({'status': 'success'})
    return JsonResponse({'status': 'error', 'message': 'Identificador Publicjud não fornecido.'}, status=400)

@staff_member_required
def relatorios_view(request):
    return render(request, 'extracao_pdf/relatorios.html')